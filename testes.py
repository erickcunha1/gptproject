import tkinter as tk
from tkinter import filedialog
from tkinter import ttk
import pandas as pd
from tkinter import messagebox
from openai import OpenAI
from docx import Document
from pathlib import Path

class InterfaceGrafica:
    def __init__(self, master):
        self.master = master
        self.master.title("Interface Gráfica 800x400")

        self.frame = tk.Frame(self.master)
        self.frame.pack(pady=20)

        self.label = tk.Label(self.frame, text="Selecione um arquivo Excel:")
        self.label.grid(row=0, column=0, padx=10, pady=10)

        self.button = tk.Button(self.frame, text="Escolher Arquivo", command=self.abrir_arquivo)
        self.button.grid(row=0, column=1, padx=10, pady=10)

        self.output_label = tk.Label(self.frame, text="")
        self.output_label.grid(row=1, column=0, columnspan=2, pady=10)

        self.search_entry = tk.Entry(self.frame, width=30)
        self.search_entry.grid(row=2, column=0, padx=10, pady=10)
        self.search_entry.bind("<KeyRelease>", self.filtrar_itens)

        self.item_combobox = ttk.Combobox(self.frame, state="readonly", width=40)
        self.item_combobox.grid(row=2, column=1, pady=10)

        self.mostrar_dados_button = tk.Button(self.frame, text="Mostrar Dados", command=self.mostrar_dados)
        self.mostrar_dados_button.grid(row=3, column=0, columnspan=2, pady=10)

        # Lista de dicionários para armazenar os dados extraídos
        self.linhas_dados = []

    def abrir_arquivo(self):
        filepath = filedialog.askopenfilename(filetypes=[("Arquivos Excel", "*.xlsx")])
        
        if filepath:
            # Tentar ler o arquivo Excel usando o pandas
            try:
                df = pd.read_excel(filepath)
                self.df_original = df.copy()  # Manter uma cópia original dos dados

                itens = df['Descrição do Item'].unique()

                self.output_label.config(text=f"Arquivo {filepath} lido com sucesso.")
                
                # Configurar os itens únicos da coluna 'Descrição do Item' no Combobox
                self.item_combobox['values'] = tuple(itens)
                self.item_combobox.set("")  # Limpar a seleção atual

                # Atualizar a lista de dicionários com os dados extraídos
                for _, linha in df.iterrows():
                    dados_linha = {coluna: linha[coluna] for coluna in df.columns}
                    self.linhas_dados.append(dados_linha)

            except Exception as e:
                self.output_label.config(text=f"Erro ao ler o arquivo:\n{e}")

    def filtrar_itens(self, event):
        # Filtrar os itens com base na entrada da barra de pesquisa
        filtro = self.search_entry.get().strip().lower()
        df_filtrado = self.df_original[self.df_original['Descrição do Item'].str.lower().str.contains(filtro)]
        itens_filtrados = df_filtrado['Descrição do Item'].unique()

        # Configurar os itens filtrados no Combobox
        self.item_combobox['values'] = tuple(itens_filtrados)
        self.item_combobox.set("")  # Limpar a seleção atual

    def mostrar_dados(self):
        item_selecionado = self.item_combobox.get()

        if item_selecionado:
            print(f"Dados para '{item_selecionado}':")

            # Encontrar a linha correspondente ao item selecionado na lista de dicionários
            for linha in self.linhas_dados:
                if linha['Descrição do Item'] == item_selecionado:
                    # Criar a lista de strings para os prompts específicos
                    lista_dados = [
                        linha['Prompt1-objeto'],
                        linha['Prompt2-Justificativa'],
                        linha['Prompt3-justificativa-quantitativo'],
                        linha['Prompt4-fundamentação legal'],
                        linha['Prompt5-detalhamento técnico'],
                        linha['Prompt6-justificativa-parcelamento'],
                        linha['Prompt7-posicionamento-conclusivo']
                    ]

                    tab_order = [
                        "Objeto",
                        "Justificativa",
                        "Justificativa do Quantitativo",
                        "Fundamentação Legal",
                        "Detalhamento Técnico",
                        "Justificativa do Parcelamento",
                        "Posicionamento Conclusivo"
                    ]

                    caminho_arquivo = Path.home() / "Desktop" / "chat.docx"
                    doc = Document()
                    client = OpenAI(api_key='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')

                    for i, prompt_valor in enumerate(lista_dados):
                        doc.add_heading(tab_order[i], level=1)
                        
                        response = client.chat.completions.create(
                            model="gpt-3.5-turbo",
                            messages=[
                                {"role": "user", "content": prompt_valor}
                            ]
                        )
                        resposta = response.choices[0].message.content

                        # Adiciona a resposta como um parágrafo na seção correspondente
                        doc.add_paragraph(resposta)

                    doc.save(caminho_arquivo)
                    messagebox.showinfo('Documento gerado!', 'Documento salvo com sucesso!')

        else:
            print("Selecione um item antes de mostrar os dados.")

if __name__ == "__main__":
    root = tk.Tk()
    root.geometry("800x400")
    app = InterfaceGrafica(root)
    root.mainloop()