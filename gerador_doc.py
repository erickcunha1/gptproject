from datetime import datetime
from pathlib import Path

from PyQt5.QtWidgets import QWidget
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
import dotenv
import openai
import os

from prompt import PromptsInfo
from prompt_menager import PromptMenager
from expressions import replace_object_description, replace_sustainability_criteria

class DocumentGenerator(QWidget):
    def __init__(self, host, user, password, database=None):
        super().__init__()
        dotenv.load_dotenv()
        openai_api_key = os.getenv('OPENAI_KEY')
        self.client = openai.OpenAI(api_key=openai_api_key)
        self.prompt_manager = PromptMenager(host, user, password, database)  # Use environment variable

    def generate_response(self, prompt):
        response = self.client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content
    
    def save_document(self, doc):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_path = Path.home() / "Desktop" / f"Generated_Documents_{timestamp}.docx"
        doc.save(file_path)
        return file_path
    
    def process_response(self, response, description, column_name_etp):
        if column_name_etp == 'object_description':
            response = replace_object_description(response).replace('$', ' ' + description + ' ')
        elif column_name_etp == 'sustainability_requirements':
            response = replace_sustainability_criteria(response).replace('$', ' ' + description + ' ')
        return response
    
    def add_image_to_header(self, doc, image_path):
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        run = header_para.add_run()
        run.add_picture(image_path, width=Inches(0.5))  # Ajuste o tamanho da imagem conforme necessário
        
        # Centralizar a imagem no cabeçalho
        header_para.alignment = 1  # 0=left, 1=center, 2=right

    def generate_etp_document(self, selected_items, objects):
        doc = Document()
        doc.add_heading("Estudo Técnico Preliminar - ETP", level=1)

        self.add_image_to_header(doc, 'gptproject/image.png')

        for obj in objects:
            unit_code = obj[:5]
            for i in range(1, 10):
                title = PromptsInfo.get_title_etp(i)
                doc.add_heading(title, level=1)
                for item in selected_items:
                    column_name = PromptsInfo.get_column_etp(i)
                    prompt_exists = self.prompt_manager.result_exists(item, column_name, 'etp', unit_code)  # Check existence in ETP table
                    if prompt_exists:
                        response = prompt_exists
                    else:
                        prompt = self.prompt_manager.search_prompt_etp(i, item, unit_code)
                        response = self.generate_response(prompt)
                        self.prompt_manager.insert_prompt(PromptsInfo.get_column_etp(i), response, item, unit_code)
                    doc.add_paragraph(response)
        self.save_document(doc)
        
    def generate_tr_document(self, selected, objects):
        doc = Document()
        doc.add_heading("TERMO DE REFERÊNCIA - TR", level=1)

        desc = None
        for obj in objects:
            unit_code = obj[:5]
            for i in range(1, 11):
                titulo = PromptsInfo.get_title_tr(i)
                doc.add_heading(text=titulo, level=1)
                column_name = PromptsInfo.get_column_tr(i)
                column_name_etp = PromptsInfo.get_column_etp(i)

                for item in selected:
                    result = None
                    if column_name_etp:
                        result = self.prompt_manager.result_exists(item, column_name_etp, 'etp', unit_code)
                        if column_name_etp == 'descricao_objeto' and result:
                            desc = result[:-1]

                    prompt_exists = self.prompt_manager.result_exists(item, column_name, 'termo_referencia', unit_code)

                    if prompt_exists and column_name_etp != 'descricao_justificativa' :
                        resposta = prompt_exists
                    elif column_name_etp == 'descricao_justificativa' and self.prompt_manager.result_exists(item, 'descricao_justificativa', 'etp', unit_code) == None and self.prompt_manager.result_exists(item, 'descricao_fundamentacao_contratacao', 'termo_referencia', unit_code) == None:
                        prompt_just = self.prompt_manager.search_prompt_etp(2, item, unit_code)
                        resposta = self.generate_response(prompt_just)
                        self.prompt_manager.insert_prompt_tr(column_name, resposta, item, unit_code)
                    else:
                        prompt = self.prompt_manager.search_prompt_tr(i, item)
                        resposta = self.generate_response(prompt)
                        self.prompt_manager.insert_prompt_tr(column_name, resposta, item, unit_code)

                    if result and column_name_etp in ['descricao_objeto', 'requisitos_contratacao']:
                        resposta = self.process_response(resposta, desc, column_name_etp)

                    doc.add_paragraph(resposta)
        self.save_document(doc)