import streamlit as st
from datetime import datetime
from pathlib import Path
from docx import Document
import dotenv
import openai
import os

from prompt import PromptsInfo
from prompt_menager import MenagerPrompt
from testclass import substituir_descricao_objeto, substituir_criterios_sustentabilidade

class GeradorDocumentos:
    def __init__(self, host, user, passwd, database=None):
        dotenv.load_dotenv()
        KEY = os.getenv('OPENAI_KEY')
        self.client = openai.OpenAI(api_key=KEY)
        self.prompt_manager = MenagerPrompt(host, user, passwd, database)  # Usar variável de ambiente

    def generate_response(self, prompt):
        response = self.client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content
    
    def save_document(self, doc):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        caminho_arquivo = Path.home() / "Desktop" / f"Documentos_Gerados_{timestamp}.docx"
        doc.save(caminho_arquivo)
        return caminho_arquivo
    
    def process_response(self, resposta, desc, column_name_etp):
        if column_name_etp == 'descricao_objeto':
            resposta = substituir_descricao_objeto(resposta).replace('$', ' ' + desc + ' ')
        elif column_name_etp == 'requisitos_contratacao':
            resposta = substituir_criterios_sustentabilidade(resposta).replace('$', ' ' + desc + ' ')
        return resposta

    def gerar_documento_etp(self, selected):
        doc = Document()
        doc.add_heading("ETP", level=1)

        for i in range(1, 9):
            titulo = PromptsInfo.get_title_etp(i)
            doc.add_heading(titulo, level=1)
            for item in selected:
                column_name = PromptsInfo.get_column_etp(i)
                prompt_exists = self.prompt_manager.result_exists(item, column_name, 'etp') # Verifica a existência na tabela ETP
                if prompt_exists:
                    resposta = prompt_exists
                else:
                    prompt = self.prompt_manager.search_prompt_etp(i, item)
                    resposta = self.generate_response(prompt)
                    self.prompt_manager.insert_prompt(PromptsInfo.get_column_etp(i), resposta, item)
                doc.add_paragraph(resposta)
        return doc
        
    def gerar_documentos_tr(self, selected):
        doc = Document()
        doc.add_heading("TERMO DE REFERÊNCIA - TR", level=1)

        desc = None
        for i in range(1, 11):
            titulo = PromptsInfo.get_title_tr(i)
            doc.add_heading(titulo, level=1)

            column_name = PromptsInfo.get_column_tr(i)
            column_name_etp = PromptsInfo.get_column_etp(i)

            for item in selected:
                result = None
                if column_name_etp:
                    result = self.prompt_manager.result_exists(item, column_name_etp, 'etp')
                    if column_name_etp == 'descricao_objeto' and result:
                        desc = result[:-1]

                prompt_exists = self.prompt_manager.result_exists(item, column_name, 'termo_referencia')

                if prompt_exists:
                    resposta = prompt_exists
                else:
                    prompt = self.prompt_manager.search_prompt_tr(i, item)
                    resposta = self.generate_response(prompt)
                    self.prompt_manager.insert_prompt_tr(column_name, resposta, item)

                if result and column_name_etp in ['descricao_objeto', 'requisitos_contratacao']:
                    resposta = self.process_response(resposta, desc, column_name_etp)

                doc.add_paragraph(resposta)
        return doc

# Tela de Login
def login():
    st.title('Login')
    st.subheader('Por favor, faça o login para continuar.')

    with st.form(key='login_form'):
        email = st.text_input('Email')
        senha = st.text_input('Senha', type='password')

        if st.form_submit_button('Login'):
            # Aqui você pode adicionar lógica para verificar as credenciais do usuário
            if email == 'erick' and senha == 'erick':
                st.success('Login bem-sucedido!')
                st.session_state.logado = True
            else:
                st.error('Email ou senha incorretos.')

    if st.button('Registrar'):
        st.session_state.tela_atual = 'registro'

# Tela de Registro
def registro():
    st.title('Registro')
    st.subheader('Por favor, preencha as informações abaixo para continuar.')

    with st.form(key='register_form'):
        nome = st.text_input('Nome')
        email = st.text_input('Email')
        senha = st.text_input('Senha', type='password')
        confirmar_senha = st.text_input('Confirmar Senha', type='password')

        if st.form_submit_button('Registrar'):
            if senha == confirmar_senha:
                st.success('Registrado com sucesso!')
                st.session_state.registrado = True
                st.session_state.tela_atual = 'login'
            else:
                st.error('As senhas não coincidem. Tente novamente.')

    if st.button('Já tem uma conta? Faça login'):
        st.session_state.tela_atual = 'login'

# Página de geração de documentos
def gerar_documentos():
    st.title('Gerador de Documentos Licitatórios')

    gerador = GeradorDocumentos('localhost', 'root', 'Erick1@3$5', 'gdl')
    cursor = gerador.prompt_manager.cursor

    st.header('Selecione os Itens')

    cursor.execute('SELECT descricao_item FROM item ORDER BY descricao_item;')
    itens1 = cursor.fetchall()
    itens = [item[0] for item in itens1]

    col1, col2 = st.columns(2)

    with col1:
        selected_items = st.multiselect('Itens', options=itens, placeholder='Selecione os itens')  # Lista de strings sem parênteses
        print(selected_items)

    with col2:
        if st.button('Gerar Documento ETP', type="primary"):
            with st.spinner('Gerando documento ETP...'):
                doc_etp = gerador.gerar_documento_etp(selected_items)
                caminho_arquivo = gerador.save_document(doc_etp)
                st.success(f'Documento ETP gerado: {caminho_arquivo}')
                with open(caminho_arquivo, 'rb') as f:
                    st.download_button(label='Baixar Documento ETP', data=f, file_name='Documento_ETP.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        if st.button('Gerar Documento TR'):
            with st.spinner('Gerando documento TR...', type="primary"):
                doc_tr = gerador.gerar_documentos_tr(selected_items)
                caminho_arquivo = gerador.save_document(doc_tr)
                st.success(f'Documento TR gerado: {caminho_arquivo}')
                with open(caminho_arquivo, 'rb') as f:
                    st.download_button(label='Baixar Documento TR', data=f, file_name='Documento_TR.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

# Configurações da aplicação Streamlit
def main():
    st.set_page_config(page_title='Gerador de Documentos Licitatórios', layout='wide')

    if 'logado' not in st.session_state:
        st.session_state.logado = False

    if 'tela_atual' not in st.session_state:
        st.session_state.tela_atual = 'login'

    if st.session_state.tela_atual == 'registro':
        registro()
    elif st.session_state.tela_atual == 'login':
        login()

    if st.session_state.logado:
        st.sidebar.title('Menu')
        selected_page = st.sidebar.radio(
            'Selecione uma opção:',
            ['Gerar Documentos', 'Sair']
        )

        if selected_page == 'Gerar Documentos':
            gerar_documentos()
        elif selected_page == 'Sair':
            st.session_state.logado = False
            st.session_state.tela_atual = 'login'

if __name__ == '__main__':
    main()