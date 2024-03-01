import sys
from PyQt5.QtWidgets import QApplication, QMainWindow, QDialog, QLabel, QLineEdit, QPushButton, QVBoxLayout, QFormLayout, QMessageBox, QDialogButtonBox, QWidget, QComboBox, QFileDialog
from PyQt5.QtGui import QPalette, QColor
from PyQt5.QtCore import Qt, pyqtSlot
import openai
from docx import Document
import pandas as pd
from datetime import datetime
from pathlib import Path
import os


class LoginDialog(QDialog):
    def __init__(self):
        super(LoginDialog, self).__init__()

        self.setWindowTitle("Login")

        layout = QVBoxLayout()

        self.username_entry = QLineEdit(self)
        self.username_entry.setPlaceholderText("Username")
        self.password_entry = QLineEdit(self)
        self.password_entry.setPlaceholderText("Password")
        self.password_entry.setEchoMode(QLineEdit.Password)

        form_layout = QFormLayout()
        form_layout.addRow("Username:", self.username_entry)
        form_layout.addRow("Password:", self.password_entry)

        buttons = QDialogButtonBox(
            QDialogButtonBox.Ok | QDialogButtonBox.Cancel,
            Qt.Horizontal, self)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)

        layout.addLayout(form_layout)
        layout.addWidget(buttons)

        self.setLayout(layout)

    def get_username_password(self):
        return self.username_entry.text(), self.password_entry.text()


class InterfaceGrafica(QMainWindow):
    def __init__(self):
        super(InterfaceGrafica, self).__init__()

        self.setWindowTitle("Gerador de Documentos Licitatórios")

        central_widget = QWidget(self)
        self.setCentralWidget(central_widget)

        layout = QVBoxLayout()

        self.label = QLabel("Selecione a origem dos dados de referência (PAC):")
        self.label.setStyleSheet("font-size: 14px; font-weight: bold; color: #2C3E50;")

        self.button = QPushButton("Escolher Arquivo", self)
        self.button.setStyleSheet("font-size: 12px; background-color: #3498DB; color: white;")
        self.button.clicked.connect(self.abrir_arquivo)

        self.output_label = QLabel("", self)
        self.output_label.setStyleSheet("font-size: 12px; color: #27AE60;")

        self.search_entry = QLineEdit(self)
        self.search_entry.setPlaceholderText("Pesquisar")
        self.search_entry.setStyleSheet("font-size: 12px;")
        self.search_entry.textChanged.connect(self.filtrar_itens)  # Connect to the textChanged signal

        self.item_combobox = QComboBox(self)
        self.item_combobox.currentIndexChanged.connect(self.atualizar_dados_selecionados)
        self.item_combobox.setStyleSheet("font-size: 12px;")

        self.mostrar_dados_button = QPushButton("Gerar Estudo Técnico Preliminar - ETP", self)
        self.mostrar_dados_button.setStyleSheet("font-size: 12px; background-color: #E74C3C; color: white;")
        self.mostrar_dados_button.clicked.connect(self.mostrar_dados)

        self.mostrar_dados_button2 = QPushButton("Gerar Termo de Referência - TR", self)
        self.mostrar_dados_button2.setStyleSheet("font-size: 12px; background-color: #E74C3C; color: white;")
        self.mostrar_dados_button2.clicked.connect(self.mostrar_dados_TR)

        self.mostrar_dados_button3 = QPushButton("Gerar Edital de Licitação", self)
        self.mostrar_dados_button3.setStyleSheet("font-size: 12px; background-color: #E74C3C; color: white;")
        self.mostrar_dados_button3.clicked.connect(self.mostrar_dados)

        self.mostrar_dados_button4 = QPushButton("Gerar Contrato", self)
        self.mostrar_dados_button4.setStyleSheet("font-size: 12px; background-color: #E74C3C; color: white;")
        self.mostrar_dados_button4.clicked.connect(self.mostrar_dados)


        layout.addWidget(self.label)
        layout.addWidget(self.button)
        layout.addWidget(self.output_label)
        layout.addWidget(self.search_entry)
        layout.addWidget(self.item_combobox)
        layout.addWidget(self.mostrar_dados_button)
        layout.addWidget(self.mostrar_dados_button2)
        layout.addWidget(self.mostrar_dados_button3)
        layout.addWidget(self.mostrar_dados_button4)

        central_widget.setLayout(layout)

        # Lista de dicionários para armazenar os dados extraídos
        self.linhas_dados = []

        # Flag para indicar se a atualização do combo box é devido à seleção manual
        self.manual_selection = True

    def abrir_arquivo(self):
        filepath, _ = QFileDialog.getOpenFileName(self, "Selecione um arquivo Excel", "", "Arquivos Excel (*.xlsx)")

        if filepath:
            try:
                df = pd.read_excel(filepath)
                self.df_original = df.copy()

                itens = df['Descrição do Item'].unique()

                self.output_label.setText(f"Arquivo {filepath} lido com sucesso.")
                self.output_label.setStyleSheet("font-size: 12px; color: #27AE60;")

                # Disconnect the signal before clearing and updating the combo box
                self.item_combobox.currentIndexChanged.disconnect(self.atualizar_dados_selecionados)
                self.item_combobox.clear()
                self.item_combobox.addItems(itens)
                self.item_combobox.setCurrentIndex(-1)
                # Reconnect the signal after updating the combo box
                self.item_combobox.currentIndexChanged.connect(self.atualizar_dados_selecionados)

                for _, linha in df.iterrows():
                    dados_linha = {coluna: linha[coluna] for coluna in df.columns}
                    self.linhas_dados.append(dados_linha)

            except Exception as e:
                self.output_label.setText(f"Erro ao ler o arquivo:\n{e}")
                self.output_label.setStyleSheet("font-size: 12px; color: #E74C3C;")

    def filtrar_itens(self):
        filtro = self.search_entry.text().strip().lower()
        df_filtrado = self.df_original[self.df_original['Descrição do Item'].str.lower().str.contains(filtro)]
        itens_filtrados = df_filtrado['Descrição do Item'].unique()

        # Disconnect the signal before clearing and updating the combo box
        self.item_combobox.currentIndexChanged.disconnect(self.atualizar_dados_selecionados)
        self.item_combobox.clear()
        self.item_combobox.addItems(itens_filtrados)
        self.item_combobox.setCurrentIndex(-1)
        # Reconnect the signal after updating the combo box
        self.item_combobox.currentIndexChanged.connect(self.atualizar_dados_selecionados)

    @pyqtSlot(int)
    def atualizar_dados_selecionados(self, index):
        # Check the flag to determine if the change is due to manual selection
        if self.manual_selection:
            item_selecionado = self.item_combobox.currentText()
            print(f"Atualizar dados para '{item_selecionado}':")
            # Implement any additional logic if needed when combo box selection changes

    def mostrar_dados(self):
        # Set the flag to False to indicate that the change is not due to manual selection
        self.manual_selection = False

        item_selecionado = self.item_combobox.currentText()

        if item_selecionado:
            print(f"Dados para '{item_selecionado}':")
            doc = Document()
            doc.add_heading("ESTUDO TÉCNICO PRELIMINAR - ETP", level=1)
            for linha in self.linhas_dados:
                if linha['Descrição do Item'] == item_selecionado:
                    lista_dados = [
                        linha['Prompt1-objeto'],
                        linha['Prompt2-Justificativa'],
                        linha['Prompt3-previsão de contratação'],
                        linha['Prompt4-requisitos da contratação'],
                        linha['Prompt5-justificativa-quantitativo'],
                        linha['Prompt6-estimativa de valor'],
                        linha['Prompt7-fundamentação legal'],
                        linha['Prompt8-justificativa-parcelamento'],
                        linha['Prompt9-posicionamento-conclusivo']
                    ]

                    tab_order = [
                        "Objeto",
                        "Justificativa da necessidade da contratação, considerado o problema a ser resolvido sob a perspectiva do interesse público",
                        "Demonstração da previsão da contratação", 
                        "Requisitos para Contratação",
                        "Estimativas e Justificativas das quantidades para a contratação",
                        "Estimativa de valor",
                        "Fundamentação Legal",
                        "Justificativa do Parcelamento",
                        "Posicionamento Conclusivo"
                    ]

                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    caminho_arquivo = Path.home() / "Desktop" / f" ETP{item_selecionado}_{timestamp}.docx"

                   
                    client = openai.OpenAI(api_key=os.environ.get('KEY')

                    for i, prompt_valor in enumerate(lista_dados):
                        doc.add_heading(tab_order[i], level=1)
                        
                        response = client.chat.completions.create(
                            model="gpt-3.5-turbo",
                            messages=[
                                {"role": "user", "content": prompt_valor}
                            ]
                        )
                        resposta = response.choices[0].message.content
                        doc.add_paragraph(resposta)

                    doc.save(caminho_arquivo)
                    QMessageBox.information(self, 'Documento gerado!', 'Documento salvo com sucesso!')

        else:
            print("Selecione um item antes de mostrar os dados.")

        # Set the flag back to True after handling the button click
        self.manual_selection = True


    def mostrar_dados_TR(self):
        # Set the flag to False to indicate that the change is not due to manual selection
        self.manual_selection = False

        item_selecionado = self.item_combobox.currentText()

        if item_selecionado:
            print(f"Dados para '{item_selecionado}':")
            doc = Document()
            doc.add_heading("TERMO DE REFERÊNCIA - TR", level=1)
            for linha in self.linhas_dados:
                if linha['Descrição do Item'] == item_selecionado:
                    lista_dados = [
                        linha['Prompt1-Termo de Referência'],
                        linha['Prompt2-TR-Fundamentação da Contratação'],
                        linha['Prompt3-TR- Descrição da solução como um todo considerado o ciclo de vida do objeto e especificação do produto'],
                        linha['Prompt4-TR- Requisitos da Contratação']


                    ]

                    tab_order = [
                        "Condições Gerais da Contratação",
                        "Fundamentação da Contratação",
                        "Descrição da solução como um todo considerado o ciclo de vida do objeto e especificação do produto",
                        "Requisitos da Contratação"
                    ]

                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    caminho_arquivo = Path.home() / "Desktop" / f" TR{item_selecionado}_{timestamp}.docx"

                   
                    client = openai.OpenAI(api_key=os.environ.get('KEY'))

                    for i, prompt_valor in enumerate(lista_dados):
                        doc.add_heading(tab_order[i], level=1)
                        
                        response = client.chat.completions.create(
                            model="gpt-3.5-turbo",
                            messages=[
                                {"role": "user", "content": prompt_valor}
                            ]
                        )
                        resposta = response.choices[0].message.content
                        doc.add_paragraph(resposta)

                    doc.save(caminho_arquivo)
                    QMessageBox.information(self, 'Documento gerado!', 'Documento salvo com sucesso!')

        else:
            print("Selecione um item antes de mostrar os dados.")

        # Set the flag back to True after handling the button click
        self.manual_selection = True



if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle('Fusion')

    login_dialog = LoginDialog()
    result = login_dialog.exec_()

    if result == QDialog.Accepted:
        username, password = login_dialog.get_username_password()

        # Perform authentication here (replace this with your own authentication logic)
        authenticated = username == "admin" and password == "admin"

        if authenticated:
            window = InterfaceGrafica()
            window.setGeometry(100, 100, 800, 400)
            window.show()
            sys.exit(app.exec_())
        else:
            QMessageBox.critical(None, 'Authentication Failed', 'Invalid username or password.')
    else:
        sys.exit()
