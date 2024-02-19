import tkinter as tk
from tkinter import ttk, simpledialog
from docx import Document
from openai import OpenAI
from pathlib import Path
import sqlite3

class SimpleGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Interface Gráfica")

        # Adicione a ordem desejada das abas
        self.tab_order = [
            "Objeto",
            "Justificativa",
            "Justificativa do Quantitativo",
            "Fundamentação Legal",
            "Detalhamento Técnico",
            "Justificativa do Parcelamento",
            "Posicionamento Conclusivo"
        ]

        # Conexão com o banco de dados SQLite
        self.connection = sqlite3.connect("prompts.db")
        self.create_table()

        self.create_tabs()

    def create_table(self):
        # Criação da tabela se não existir
        cursor = self.connection.cursor()
        cursor.execute('''CREATE TABLE IF NOT EXISTS prompts
                          (id INTEGER PRIMARY KEY AUTOINCREMENT,
                           topic TEXT,
                           prompt TEXT)''')
        self.connection.commit()

    def create_tabs(self):
        self.notebook = ttk.Notebook(self.root)

        for tab_name in self.tab_order:
            tab_frame = ttk.Frame(self.notebook)
            self.create_listbox(tab_frame, tab_name)
            self.notebook.add(tab_frame, text=tab_name)

        self.notebook.pack(expand=True, fill="both")

        generate_btn = ttk.Button(self.root, text="Gerar Documento", command=self.generate_document)
        generate_btn.pack(pady=10)

    def create_listbox(self, frame, prompt_key):
        prompts_listbox = tk.Listbox(frame, selectmode=tk.SINGLE, width=50, height=10)
        prompts_listbox.pack(pady=10)

        add_btn = ttk.Button(frame, text="Adicionar", command=lambda: self.add_prompt(prompt_key, prompts_listbox))
        add_btn.pack(pady=5)

        edit_btn = ttk.Button(frame, text="Editar", command=lambda: self.edit_prompt(prompt_key, prompts_listbox))
        edit_btn.pack(pady=5)

        delete_btn = ttk.Button(frame, text="Excluir", command=lambda: self.delete_prompt(prompt_key, prompts_listbox))
        delete_btn.pack(pady=5)

        # Carregar prompts do banco de dados
        prompts = self.get_prompts_from_db(prompt_key)
        for prompt in prompts:
            prompts_listbox.insert(tk.END, prompt)

    def add_prompt(self, prompt_key, listbox):
        prompt = simpledialog.askstring("Adicionar Prompt", f"Digite o prompt para {prompt_key}:", parent=self.root)
        if prompt:
            listbox.insert(tk.END, prompt)
            self.save_prompt_to_db(prompt_key, prompt)

    def edit_prompt(self, prompt_key, listbox):
        selected_index = listbox.curselection()
        if selected_index:
            selected_prompt = listbox.get(selected_index)
            new_prompt = simpledialog.askstring("Editar Prompt", f"Edite o prompt para {prompt_key}:", initialvalue=selected_prompt)

            if new_prompt:
                listbox.delete(selected_index)
                listbox.insert(selected_index, new_prompt)
                self.update_prompt_in_db(prompt_key, selected_prompt, new_prompt)

    def delete_prompt(self, prompt_key, listbox):
        selected_index = listbox.curselection()
        if selected_index:
            deleted_prompt = listbox.get(selected_index)
            listbox.delete(selected_index)
            self.delete_prompt_in_db(prompt_key, deleted_prompt)

    def save_prompt_to_db(self, topic, prompt):
        cursor = self.connection.cursor()
        cursor.execute("INSERT INTO prompts (topic, prompt) VALUES (?, ?)", (topic, prompt))
        self.connection.commit()

    def update_prompt_in_db(self, topic, old_prompt, new_prompt):
        cursor = self.connection.cursor()
        cursor.execute("UPDATE prompts SET prompt = ? WHERE topic = ? AND prompt = ?", (new_prompt, topic, old_prompt))
        self.connection.commit()

    def delete_prompt_in_db(self, topic, prompt):
        cursor = self.connection.cursor()
        cursor.execute("DELETE FROM prompts WHERE topic = ? AND prompt = ?", (topic, prompt))
        self.connection.commit()

    def get_prompts_from_db(self, topic):
        cursor = self.connection.cursor()
        cursor.execute("SELECT prompt FROM prompts WHERE topic = ?", (topic,))
        return [row[0] for row in cursor.fetchall()]

    def generate_document(self):
        caminho_arquivo = Path.home() / "Desktop" / "chat.docx"
        doc = Document()

        for topico, prompt_key in zip(self.tab_order, self.tab_order):
            prompts = self.get_prompts_from_db(prompt_key)
            client = OpenAI(api_key='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')

            # Adiciona uma seção no documento para cada topico
            doc.add_heading(topico, level=1)

            for prompt in prompts:
                response = client.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "user", "content": prompt}
                    ]
                )
                resposta = response.choices[0].message.content

                # Adiciona a resposta como um parágrafo na seção correspondente
                doc.add_paragraph(resposta)

        doc.save(caminho_arquivo)
        print("Conteúdo modificado foi salvo com sucesso.")

    def __del__(self):
        # Fechar a conexão com o banco de dados quando a instância for destruída
        self.connection.close()

if __name__ == "__main__":
    root = tk.Tk()
    app = SimpleGUI(root)
    root.mainloop()