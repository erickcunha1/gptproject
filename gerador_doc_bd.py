from PyQt5.QtWidgets import QMainWindow, QMessageBox
import openai
from docx import Document
from datetime import datetime
from pathlib import Path
import os
import mysql.connector
from mysql.connector import Error
from prompts import PromptsInfo



try:
    connection = mysql.connector.connect(
        host='127.0.0.1',
        database='bd_gdl',
        user='root',
        password='root@2024'
    )
    print("conectou")
    breakpoint
    if connection.is_connected():
        db_Info = connection.get_server_info()
        print("Conectado ao servidor MySQL versão ", db_Info)
        cursor = connection.cursor()
        cursor.execute("select database();")
        record = cursor.fetchone()
        print("Você está conectado ao banco de dados: ", record)

except Error as e:
    print("Erro ao conectar ao MySQL", e)

class InterfaceGrafica(QMainWindow):
        
    def abrir_arquivo(self):
        
        try:
            prompts_info = PromptsInfo(self.host, self.user, self.passwd, self.database)
            prompts_info.cursor.execute('SELECT descricao_item FROM item order by descricao_item;')
            itens = prompts_info.cursor.fetchall()
            self.df_original = itens
            for item in itens:
                valor = item[0]  # Obter o valor da coluna
                valor_sem_parenteses = valor.strip('()')  # Remover os parênteses
                self.item_list.addItem(valor_sem_parenteses)  # Adicionar o valor à lista na interface gráfica
        except Exception as e:
            QMessageBox.warning(self, 'Erro ao abrir arquivo', f'Ocorreu um erro ao abrir o arquivo: {str(e)}')


    def filtrar_itens(self):
        filtro = self.search_entry.text().strip().lower()
        cursor = connection.cursor()
        comando_sql = f"SELECT * FROM item WHERE LOWER(descricao_item) LIKE '%{filtro}%' order by descricao_item"
        cursor.execute(comando_sql)
        dados_lidos = cursor.fetchall()
        itens_filtrados = [item[1] for item in dados_lidos]  # Assumindo que a descrição está na segunda coluna
        itens_selecionados = [item.text() for item in self.item_list.selectedItems()]

        self.item_list.clear()
        self.item_list.addItems(itens_filtrados)

            # Restaure a seleção dos itens previamente selecionados
        for index in range(self.item_list.count()):
                item = self.item_list.item(index)
                if item.text() in itens_selecionados:
                    item.setSelected(True)


    def mostrar_dados(self):
        selected_items = [item.text() for item in self.item_list.selectedItems()]

        if selected_items:
            doc = Document()
            doc.add_heading("ESTUDO TÉCNICO PRELIMINAR - ETP", level=1)  # Adiciona um cabeçalho geral
            n = 0
            for i in range(1, 10): 
                n= i+2
                titulo = self.get_titulo_name(i)  # Obtem o título correspondente
                doc.add_heading(titulo, level=1)  # Adiciona o título

                for item_selecionado in selected_items:

                    cursor = connection.cursor()
                    query = "SELECT distinct item.cod_item, item.descricao_item, prompt_etp.cod_item, prompt_etp.descricao_objeto, prompt_etp.descricao_justificativa, prompt_etp.descricao_previsao_contratacao, prompt_etp.requisitos_contratacao, prompt_etp.descricao_justificativa_quantidade, prompt_etp.prompt_estimativa_valor, prompt_etp.prompt_fundamentacao_legal, prompt_etp.prompt_justificativa_parcelamento, prompt_etp.prompt_posicionamento_conclusivo FROM item JOIN prompt_etp ON item.cod_item = prompt_etp.cod_item WHERE descricao_item = %s"
                    cursor.execute(query, (item_selecionado,))
                    
                   # Busca o resultado
                    record = cursor.fetchone()
                            
                    if record:
                        print("Código do Item: ", record[0], "Descrição: ", record[n])
                        print(n)
                    else:
                        print("Item não encontrado.")                   
                   
                    if record[1] == item_selecionado:
                            prompt_valor = record[n]  # Obtem o va
                            client = openai.OpenAI(api_key='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')
                            response = client.chat.completions.create(
                                model="gpt-4-turbo-2024-04-09",
                                messages=[{"role": "user", "content": prompt_valor}])
                            resposta = response.choices[0].message.content
                            doc.add_paragraph(resposta)
                            n= i+2
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            caminho_arquivo = Path.home() / "Desktop" / f"Documentos_Gerados_{timestamp}.docx"
            doc.save(caminho_arquivo)
            QMessageBox.information(self, 'Documentos gerados!', f'Documento salvo com sucesso em: {caminho_arquivo}')
        else:
            print("Selecione um item antes de mostrar os dados.")

            
    def mostrar_dados_TR(self):
        selected_items = [item.text() for item in self.item_list.selectedItems()]

        if selected_items:
            doc = Document()
            doc.add_heading("TERMO DE REFERÊNCIA - TR", level=1)
            n = 0
            for i in range(1, 11):
                n= i+2
                titulo = self.get_titulo_tr(i)
                doc.add_heading(text=titulo, level=1)

                for item_selecionado in selected_items:
                    cursor = connection.cursor()

                    query = "SELECT distinct item.cod_item, item.descricao_item, prompt_tr.cod_item, prompt_tr.descricao_objeto_tr, prompt_tr.descricao_fundamentacao_contratacao, prompt_tr.descricao_solucao, prompt_tr.requisitos_contratacao, prompt_tr.modelo_execucao_objeto, prompt_tr.modelo_gestao_contrato, prompt_tr.criterios_medicao_pagamento, prompt_tr.criterios_selecao_fornecedor, prompt_tr.estimativas_valor, prompt_tr.adequacao_orcamentaria FROM item JOIN prompt_tr ON item.cod_item = prompt_tr.cod_item WHERE descricao_item = %s"

                    cursor.execute(query, (item_selecionado,))

                   # Busca o resultado
                    record = cursor.fetchone()
                            
                    if record:
                        print(n)
                    else:
                        print("Item não encontrado.")                   
                
                    if record[1] == item_selecionado:
                            prompt_valor = record[n]  # Obtem o valor do prompt
                            print(prompt_valor)
                            client = openai.OpenAI(api_key=os.environ.get('KEY'))

                            response = client.chat.completions.create(
                                model="gpt-4-turbo-2024-04-09",
                                messages=[
                                    {"role": "system", "content": "Você é um especialista em elaboração de licitações na administração pública e deve gerar as respostas com base nos detalhes dos prompts. Detalhe cada resposta de acordo com o promt recebido, sem egolir linhas"},                               
                                    {"role": "user", "content": prompt_valor}])
                            resposta = response.choices[0].message.content
                            print(resposta)
                            doc.add_paragraph(resposta)
                            n= i+2

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            caminho_arquivo = Path.home() / "Desktop" / f"Documentos_Gerados_{timestamp}.docx"

            doc.save(caminho_arquivo)
            QMessageBox.information(self, 'Documentos gerados!', f'Documento salvo com sucesso em: {caminho_arquivo}')
        else:
            print("Selecione um item antes de mostrar os dados.")

